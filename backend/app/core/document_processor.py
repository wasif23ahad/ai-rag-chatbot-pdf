"""
Document Processor — PDF and DOCX parsing, text cleaning, and chunking.
Produces LangChain Document objects with deterministic chunk metadata.
"""

import re
import unicodedata
from typing import List, Tuple

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document


class DocumentProcessor:
    """
    Routes files to the appropriate extractor, cleans text, and splits into chunks.
    Each chunk receives deterministic metadata: chunk_id, page, source, char_start.
    """

    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 100) -> None:
        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
            length_function=len,
            is_separator_regex=False,
        )

    def process(self, file_path: str, file_type: str) -> List[Document]:
        """
        Full pipeline: extract → clean → split → attach metadata.

        Args:
            file_path: Absolute path to the temp file on disk.
            file_type: File extension including dot — '.pdf' or '.docx'.

        Returns:
            List[Document] with .page_content and .metadata populated.

        Raises:
            ValueError: If file_type is unsupported or document has no text.
        """
        ext = file_type.lower()

        if ext == ".pdf":
            raw_pages = self._extract_pdf(file_path)
        elif ext == ".docx":
            raw_pages = self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: '{ext}'. Allowed: .pdf, .docx")

        if not raw_pages:
            raise ValueError("Document contains no extractable text.")

        source_name = file_path.split("/")[-1].split("\\")[-1]
        return self._chunk_pages(raw_pages, source=source_name)

    def _extract_pdf(self, path: str) -> List[Tuple[str, int]]:
        """
        Extract text from a PDF file, preserving page numbers (1-indexed).

        Returns:
            List of (page_text, page_number) tuples — empty pages are skipped.

        Raises:
            ValueError: If no extractable text is found (e.g. scanned/image PDF).
        """
        from pypdf import PdfReader  # imported here to keep module import fast

        reader = PdfReader(path)
        pages: List[Tuple[str, int]] = []

        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            cleaned = self._clean_text(text)
            if cleaned:
                pages.append((cleaned, page_num))

        if not pages:
            raise ValueError(
                "No extractable text found in PDF. "
                "The file may be image-based (scanned). "
                "Use a PDF with selectable text."
            )

        return pages

    def _extract_docx(self, path: str) -> List[Tuple[str, int]]:
        """
        Extract paragraphs and table cells from a DOCX file.
        DOCX has no native page numbers — paragraph index is used instead.

        Returns:
            List of (text, index) tuples — empty paragraphs are skipped.

        Raises:
            ValueError: If no extractable text is found.
        """
        from docx import Document as DocxDocument  # imported here to keep module import fast

        doc = DocxDocument(path)
        items: List[Tuple[str, int]] = []
        idx = 1

        for para in doc.paragraphs:
            text = self._clean_text(para.text)
            if text:
                items.append((text, idx))
                idx += 1

        # Tables are not in doc.paragraphs — iterate separately
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = self._clean_text(cell.text)
                    if text:
                        items.append((text, idx))
                        idx += 1

        if not items:
            raise ValueError("No extractable text found in DOCX.")

        return items

    def _clean_text(self, text: str) -> str:
        """
        Normalize and clean raw extracted text.
        - Normalize unicode (smart quotes, ligatures → ASCII equivalents)
        - Strip null bytes and control characters (preserve \\n and \\t)
        - Collapse multiple spaces/tabs to a single space
        - Collapse 3+ newlines to double newline (preserve paragraph breaks)
        - Remove purely decorative lines (---, ===, ...)
        """
        if not text:
            return ""

        # Unicode normalization: ligatures, smart quotes, etc.
        text = unicodedata.normalize("NFKC", text)

        # Strip null bytes and non-printable control chars (keep \n \t)
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

        # Collapse horizontal whitespace
        text = re.sub(r"[ \t]+", " ", text)

        # Collapse excessive newlines
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Drop purely decorative lines
        lines = [
            line for line in text.split("\n")
            if not re.match(r"^[\s\-=_\*\.]{3,}$", line.strip())
        ]
        text = "\n".join(lines)

        return text.strip()

    def _chunk_pages(
        self,
        pages: List[Tuple[str, int]],
        source: str,
    ) -> List[Document]:
        """
        Split page texts into overlapping chunks and attach deterministic metadata.

        chunk_id format: 'chunk_0000', 'chunk_0001', ...
        char_start: character offset within the page text (not the whole document).
        """
        chunks: List[Document] = []
        chunk_counter = 0

        for page_text, page_num in pages:
            raw_chunks = self._splitter.split_text(page_text)
            char_pos = 0

            for chunk_text in raw_chunks:
                doc = Document(
                    page_content=chunk_text,
                    metadata={
                        "chunk_id": f"chunk_{chunk_counter:04d}",
                        "page": page_num,
                        "source": source,
                        "char_start": char_pos,
                    },
                )
                chunks.append(doc)
                char_pos += len(chunk_text)
                chunk_counter += 1

        if not chunks:
            raise ValueError("Text splitting produced zero chunks.")

        return chunks
